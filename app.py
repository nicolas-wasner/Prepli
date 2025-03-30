from flask import Flask, redirect, render_template, request, send_file, flash
from docx import Document
from flask_sqlalchemy import SQLAlchemy
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
import os
import io
from dotenv import load_dotenv

load_dotenv()  # charge les variables depuis un fichier .env

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv("DATABASE_URL")
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    "pool_pre_ping": True,
    "pool_recycle": 280
}

db = SQLAlchemy(app)



# Définition du modèle Fiche
class Fiche(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    domaine = db.Column(db.String(200))
    duree = db.Column(db.String(100))
    niveau = db.Column(db.String(100))
    sequence = db.Column(db.String(200))
    seance = db.Column(db.String(200))
    competences_visees = db.Column(db.Text)
    competences_scccc = db.Column(db.Text)
    objectifs = db.Column(db.Text)
    afc = db.Column(db.String(200))
    prerequis = db.Column(db.Text)
    phases = db.Column(db.Text)
    deroulement = db.Column(db.Text)
    consigne = db.Column(db.Text)
    nom_enseignant = db.Column(db.String(200))
    eleve = db.Column(db.Text)
    differenciation = db.Column(db.Text)
    materiel = db.Column(db.Text)
    sequence_id = db.Column(db.Integer, db.ForeignKey('sequence.id'), nullable=True)
    modalites_evaluation = db.Column(db.Text)
    bilan = db.Column(db.Text)
    prolongements = db.Column(db.Text)
    remediations = db.Column(db.Text)


class Sequence(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    titre = db.Column(db.String(200), nullable=False)
    seances = db.relationship('Fiche', backref='sequence_obj', lazy=True)


# Créer la base de données (si elle n'existe pas déjà)
with app.app_context():
    db.create_all()

@app.route('/creer')
def creer_fiche():
    return render_template('creer_fiche.html')


@app.route('/')
def dashboard():
    return render_template('index.html')


# Route pour enregistrer la fiche dans la base
@app.route('/submit', methods=['POST'])
def submit():
    fiche = Fiche(
        domaine=request.form.get("domaine"),
        duree=request.form.get("duree"),
        niveau=request.form.get("niveau"),
        sequence=request.form.get("sequence"),
        seance=request.form.get("seance"),
        competences_visees=request.form.get("competences_visees"),
        competences_scccc=request.form.get("competences_scccc"),
        objectifs=request.form.get("objectifs"),
        afc=request.form.get("afc"),
        prerequis=request.form.get("prerequis"),
        phases=request.form.get("phases"),
        deroulement=request.form.get("deroulement"),
        consigne=request.form.get("consigne"),
        nom_enseignant=request.form.get("enseignant"),
        eleve=request.form.get("eleve"),
        differenciation=request.form.get("differenciation"),
        materiel=request.form.get("materiel"),
        


    )
    db.session.add(fiche)
    db.session.commit()
    flash("Fiche créée avec succès ✅", "success")
    return render_template("confirmation.html", message="Fiche créée avec succès ✅")



@app.route('/fiches')
def fiches():
    query = Fiche.query

    # Récupération des filtres
    recherche = request.args.get('recherche')
    niveau = request.args.get('niveau')
    domaine = request.args.get('domaine')

    # Appliquer les filtres si fournis
    if recherche:
        query = query.filter(
            Fiche.seance.ilike(f"%{recherche}%") |
            Fiche.sequence.ilike(f"%{recherche}%") |
            Fiche.domaine.ilike(f"%{recherche}%")
        )
    if niveau:
        query = query.filter(Fiche.niveau == niveau)
    if domaine:
        query = query.filter(Fiche.domaine == domaine)

    fiches = query.order_by(Fiche.id.desc()).all()
    niveaux = db.session.query(Fiche.niveau).distinct().all()
    domaines = db.session.query(Fiche.domaine).distinct().all()

    return render_template('fiches.html', fiches=fiches, niveaux=niveaux, domaines=domaines)

@app.route('/sequences')
def liste_sequences():
    sequences = Sequence.query.all()
    return render_template('liste_sequences.html', sequences=sequences)

@app.route('/sequences/modifier/<int:id>', methods=['GET', 'POST'])
def modifier_sequence(id):
    sequence = Sequence.query.get_or_404(id)
    toutes_les_fiches = Fiche.query.all()

    if request.method == 'POST':
        sequence.titre = request.form.get('titre')
        ids_seances = request.form.getlist('seances')

        # Dissocier toutes les séances actuelles
        for fiche in sequence.seances:
            fiche.sequence_id = None

        # Réassocier les nouvelles
        for id_seance in ids_seances:
            fiche = Fiche.query.get(int(id_seance))
            fiche.sequence_id = sequence.id

        db.session.commit()
        return redirect('/sequences')

    return render_template('modifier_sequence.html', sequence=sequence, fiches=toutes_les_fiches)


@app.route('/sequences/creer', methods=['GET', 'POST'])
def creer_sequence():
    fiches = Fiche.query.all()

    if request.method == 'POST':
        titre = request.form.get('titre')
        ids_seances = request.form.getlist('seances')

        sequence = Sequence(titre=titre)
        db.session.add(sequence)
        db.session.commit()

        for id_seance in ids_seances:
            fiche = Fiche.query.get(int(id_seance))
            fiche.sequence_id = sequence.id
        db.session.commit()

        return redirect('/sequences')

    return render_template('creer_sequence.html', fiches=fiches)



@app.route('/fiche/<int:id>')
def fiche_detail(id):
    fiche = Fiche.query.get_or_404(id)
    return render_template('fiche_detail.html', fiche=fiche)

@app.route('/modifier/<int:id>', methods=['GET', 'POST'])
def modifier_fiche(id):
    fiche = Fiche.query.get_or_404(id)

    if request.method == 'POST':
        fiche.domaine = request.form.get("domaine")
        fiche.duree = request.form.get("duree")
        fiche.niveau = request.form.get("niveau")
        fiche.sequence = request.form.get("sequence")
        fiche.seance = request.form.get("seance")
        fiche.competences_visees = request.form.get("competences_visees")
        fiche.competences_scccc = request.form.get("competences_scccc")
        fiche.objectifs = request.form.get("objectifs")
        fiche.afc = request.form.get("afc")
        fiche.prerequis = request.form.get("prerequis")
        fiche.phases = request.form.get("phases")
        fiche.deroulement = request.form.get("deroulement")
        fiche.consigne = request.form.get("consigne")
        fiche.enseignant = request.form.get("enseignant")
        fiche.eleve = request.form.get("eleve")
        fiche.differenciation = request.form.get("differenciation")
        fiche.materiel = request.form.get("materiel")
        fiche.nom_enseignant = request.form.get("nom_enseignant")
        fiche.modalites_evaluation = request.form.get("modalites_evaluation")
        fiche.bilan = request.form.get("bilan")
        fiche.prolongements = request.form.get("prolongements")
        fiche.remediations = request.form.get("remediations")

        
        db.session.commit()
        flash("Fiche créée avec succès ✅", "success")
        return render_template("confirmation.html", message="Fiche modifié avec succès ✅")

    return render_template("modifier_fiche.html", fiche=fiche)

@app.route('/supprimer/<int:id>')
def supprimer_fiche(id):
    fiche = Fiche.query.get_or_404(id)
    db.session.delete(fiche)
    db.session.commit()
    return f"Fiche #{id} supprimée. <a href='/fiches'>Retour</a>"

from docx import Document
from flask import send_file
import io


@app.route('/exporter/<int:id>')
def exporter_fiche_finale(id):
    fiche = Fiche.query.get_or_404(id)
    doc = Document()

    # Format paysage
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Style global
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Fonction utils
    def add_row(table, items):
        row = table.add_row().cells
        for i, (label, value) in enumerate(items):
            p = row[i].paragraphs[0]
            run = p.add_run(label)
            run.bold = True
            p.add_run(f" {value or ''}")

    def add_single_cell_row(table, label, content=""):
        cell = table.add_row().cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(label + "\n")
        run.bold = True
        p.add_run(content or "")

    # --- 1. Tableau principal : en-tête ---
    table1 = doc.add_table(rows=0, cols=3)
    table1.style = 'Table Grid'
    add_row(table1, [
        ("Domaine d’apprentissage :", fiche.domaine),
        ("Niveau :", fiche.niveau),
        ("Durée totale de la séance :", fiche.duree)
    ])
    add_row(table1, [
        ("Place de la séance dans la séquence :", "/"),
        ("Titre de la séquence :", fiche.sequence),
        ("Titre de la séance :", fiche.seance)
    ])
    add_row(table1, [
        ("Objectif(s) visé(s) :", fiche.objectifs),
        ("Compétence(s) visée(s) :", fiche.competences_visees),
        ("AFC :", fiche.afc)
    ])
    # Ligne unique pour Prérequis
    row = table1.add_row().cells
    cell = row[0]
    p = cell.paragraphs[0]
    run = p.add_run("Prérequis :")
    run.bold = True
    p.add_run(f" {fiche.prerequis or ''}")
    cell.merge(row[1]).merge(row[2])

    # --- 2. Tableau central à 7 colonnes ---
    table2 = doc.add_table(rows=1, cols=7)
    table2.style = 'Table Grid'
    headers = [
        "Phase et durée", "Déroulement", "Consigne",
        "Rôle de l’enseignant", "Rôle de l’élève",
        "Différenciation", "Matériel"
    ]
    for i, h in enumerate(headers):
        table2.cell(0, i).text = h

    row = table2.add_row().cells
    row[0].text = fiche.phases or ''
    row[1].text = fiche.deroulement or ''
    row[2].text = fiche.consigne or ''
    row[3].text = fiche.nom_enseignant or ''
    row[4].text = fiche.eleve or ''
    row[5].text = fiche.differenciation or ''
    row[6].text = fiche.materiel or ''

    # --- 3. Bas de fiche : 1 colonne par ligne ---
    table3 = doc.add_table(rows=0, cols=1)
    table3.style = 'Table Grid'
    add_single_cell_row(table3, "Modalités d’évaluation :", fiche.modalites_evaluation)
    add_single_cell_row(table3, "Bilan pédagogique et didactique :", fiche.bilan)
    add_single_cell_row(table3, "Prolongement(s) possible(s) :", fiche.prolongements)
    add_single_cell_row(table3, "Remédiation(s) éventuelle(s) :", fiche.remediations)
    add_single_cell_row(table3, "Nom / Prénom de l’enseignant :", fiche.nom_enseignant)

    # Exporter dans un buffer mémoire
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"fiche_{fiche.id}_{fiche.seance}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# Lancer l'application
if __name__ == "__main__":
    app.run()
